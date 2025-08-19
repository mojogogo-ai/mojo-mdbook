#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Document Converter
将Mojo GoGo项目的Markdown章节转换为Word文档
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_element(name):
    """创建XML元素"""
    return OxmlElement(name)

def create_attribute(element, name, value):
    """创建XML属性"""
    element.set(qn(name), value)

def add_page_break_before_paragraph(paragraph):
    """在段落前添加分页符"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    tag = run._r
    tag.append(create_element('w:br'))
    tag.append(create_element('w:br'))
    tag.append(create_element('w:br'))

def convert_markdown_to_word(md_file_path, output_dir):
    """将Markdown文件转换为Word文档"""
    
    # 读取Markdown文件
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 处理Markdown内容
    lines = content.split('\n')
    current_list_level = 0
    in_code_block = False
    in_style_block = False
    
    for line in lines:
        line = line.strip()
        
        # 跳过样式块
        if line.startswith('<style>'):
            in_style_block = True
            continue
        if line.startswith('</style>'):
            in_style_block = False
            continue
        if in_style_block:
            continue
        
        # 跳过HTML标签行（只保留图片标签）
        if '<div' in line or '</div>' in line:
            continue
        if '<span' in line or '</span>' in line:
            continue
        if '<p' in line or '</p>' in line:
            continue
        if '<br' in line or '<br/>' in line:
            continue
        if '<strong' in line or '</strong>' in line:
            continue
        if '<em' in line or '</em>' in line:
            continue
        if '<a href' in line or '</a>' in line:
            continue
        
        # 跳过空行
        if not line:
            continue
        
        # 处理标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()
            
            if level == 1:
                # 主标题
                heading = doc.add_heading(title_text, level=0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 添加分页符
                add_page_break_before_paragraph(heading)
            else:
                # 子标题
                doc.add_heading(title_text, level=level-1)
        
        # 处理图片 - Markdown格式
        elif line.startswith('!['):
            # 提取图片路径和alt文本
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text = match.group(1)
                img_path = match.group(2)
                
                # 构建完整的图片路径
                if img_path.startswith('../assets/'):
                    img_path = img_path.replace('../assets/', 'src/assets/')
                elif img_path.startswith('./assets/'):
                    img_path = img_path.replace('./assets/', 'src/assets/')
                elif img_path.startswith('../'):
                    img_path = img_path.replace('../', 'src/')
                elif img_path.startswith('./'):
                    img_path = img_path.replace('./', 'src/')
                elif not img_path.startswith('src/'):
                    img_path = f"src/assets/{img_path}"
                
                # 检查图片是否存在
                if os.path.exists(img_path):
                    try:
                        # 添加图片
                        doc.add_picture(img_path, width=Inches(5))
                        # 添加图片说明
                        if alt_text:
                            caption = doc.add_paragraph(alt_text)
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"无法添加图片 {img_path}: {e}")
                        # 添加图片占位符
                        doc.add_paragraph(f"[图片: {alt_text}]")
                else:
                    # 图片不存在时添加占位符
                    doc.add_paragraph(f"[图片: {alt_text}] - 文件路径: {img_path}")
        
        # 处理HTML格式的图片
        elif '<img' in line and 'src=' in line:
            # 提取图片路径和alt文本
            src_match = re.search(r'src=["\']([^"\']+)["\']', line)
            alt_match = re.search(r'alt=["\']([^"\']*)["\']', line)
            
            if src_match:
                img_path = src_match.group(1)
                alt_text = alt_match.group(1) if alt_match else ""
                
                # 构建完整的图片路径
                if img_path.startswith('../assets/'):
                    img_path = img_path.replace('../assets/', 'src/assets/')
                elif img_path.startswith('./assets/'):
                    img_path = img_path.replace('./assets/', 'src/assets/')
                elif img_path.startswith('../'):
                    img_path = img_path.replace('../', 'src/')
                elif img_path.startswith('./'):
                    img_path = img_path.replace('./', 'src/')
                elif not img_path.startswith('src/'):
                    img_path = f"src/assets/{img_path}"
                
                # 检查图片是否存在
                if os.path.exists(img_path):
                    try:
                        # 添加图片
                        doc.add_picture(img_path, width=Inches(5))
                        # 添加图片说明
                        if alt_text:
                            caption = doc.add_paragraph(alt_text)
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except Exception as e:
                        print(f"无法添加图片 {img_path}: {e}")
                        # 添加图片占位符
                        doc.add_paragraph(f"[图片: {alt_text}]")
                else:
                    # 图片不存在时添加占位符
                    doc.add_paragraph(f"[图片: {alt_text}] - 文件路径: {img_path}")
        
        # 处理链接
        elif '[' in line and '](' in line and ')' in line:
            # 提取链接文本和URL
            match = re.match(r'\[(.*?)\]\((.*?)\)', line)
            if match:
                link_text = match.group(1)
                url = match.group(2)
                
                # 创建带链接的段落
                p = doc.add_paragraph()
                run = p.add_run(link_text)
                run.font.color.rgb = None  # 默认颜色
                p.add_run(f" ({url})")
        
        # 处理列表项
        elif line.startswith('- ') or line.startswith('* '):
            list_text = line[2:].strip()
            p = doc.add_paragraph(list_text, style='List Bullet')
        
        # 处理数字列表
        elif re.match(r'^\d+\.', line):
            list_text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(list_text, style='List Number')
        
        # 处理代码块
        elif line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                # 开始代码块
                doc.add_paragraph("代码块:")
            continue
        
        # 处理代码块内容
        elif in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            run = p.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        # 处理引用块
        elif line.startswith('> '):
            quote_text = line[2:].strip()
            p = doc.add_paragraph(quote_text)
            p.style = 'Quote'
        
        # 处理警告和提示框
        elif line.startswith('⚠️') or line.startswith('💡') or line.startswith('📋') or line.startswith('⏱️') or line.startswith('🎉'):
            p = doc.add_paragraph(line)
            p.style = 'Intense Quote'
        
        # 处理普通段落
        else:
            # 检查是否包含粗体或斜体
            if '**' in line or '*' in line:
                # 处理粗体和斜体
                p = doc.add_paragraph()
                current_pos = 0
                
                # 查找所有格式标记
                format_markers = []
                for match in re.finditer(r'\*\*([^*]+)\*\*|\*([^*]+)\*', line):
                    format_markers.append((match.start(), match.end(), match.group(0), match.group(1) or match.group(2)))
                
                if format_markers:
                    for i, (start, end, marker, text) in enumerate(format_markers):
                        # 添加格式标记前的文本
                        if start > current_pos:
                            p.add_run(line[current_pos:start])
                        
                        # 添加格式化的文本
                        run = p.add_run(text)
                        if marker.startswith('**'):
                            run.bold = True
                        else:
                            run.italic = True
                        
                        current_pos = end
                    
                    # 添加剩余的文本
                    if current_pos < len(line):
                        p.add_run(line[current_pos:])
                else:
                    p.add_run(line)
            else:
                # 普通段落
                doc.add_paragraph(line)
    
    # 保存文档
    filename = os.path.basename(md_file_path).replace('.md', '.docx')
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    print(f"已生成: {output_path}")
    
    return output_path

def main():
    """主函数"""
    # 创建输出目录
    output_dir = "word_documents"
    os.makedirs(output_dir, exist_ok=True)
    
    # 定义要转换的章节
    chapters = [
        "src/chapters/user-guide.md",
        "src/chapters/roam-points.md", 
        "src/chapters/create-bot.md",
        "src/chapters/earn-points.md"
    ]
    
    print("开始转换Markdown章节为Word文档...")
    print("=" * 50)
    
    converted_files = []
    
    for chapter in chapters:
        if os.path.exists(chapter):
            try:
                output_file = convert_markdown_to_word(chapter, output_dir)
                converted_files.append(output_file)
            except Exception as e:
                print(f"转换 {chapter} 时出错: {e}")
        else:
            print(f"文件不存在: {chapter}")
    
    print("=" * 50)
    print(f"转换完成! 共生成 {len(converted_files)} 个Word文档")
    print(f"输出目录: {os.path.abspath(output_dir)}")
    
    # 列出生成的文件
    print("\n生成的文件:")
    for file in converted_files:
        print(f"  - {os.path.basename(file)}")

if __name__ == "__main__":
    main()
